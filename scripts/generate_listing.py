import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_program_listing():
    # Initialize document
    doc = Document()
    doc.add_heading('Листинг программы "Аналитика УИ ОГПЗ"', 0)

    # List of files to include with their descriptions
    files_to_process = [
        {
            'path': 'analytics_ui/excel_merger.py',
            'desc': 'Главный модуль приложения. Содержит реализацию графического интерфейса (GUI) на базе Tkinter, а также основную бизнес-логику: чтение Excel-файлов, их объединение, фильтрацию данных, привязку к правилам и применение условного форматирования (цвета, границы, стрелки трендов) в итоговом отчете.'
        },
        {
            'path': 'analytics_ui/post_install.py',
            'desc': 'Скрипт пост-установки. Отвечает за интеграцию приложения в рабочее окружение Linux (RedOS). Создает файлы ярлыков (.desktop) на рабочем столе и в системном меню приложений для удобного запуска программы пользователем.'
        },
        {
            'path': 'setup.py',
            'desc': 'Файл конфигурации пакета Python. Определяет метаданные проекта (имя, версия, автор), зависимости (pandas, openpyxl и др.), точки входа (консольные команды analytics-ui) и правила сборки дистрибутива (.whl).'
        }
    ]
    
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    for item in files_to_process:
        relative_path = item['path']
        description = item['desc']
        file_path = os.path.join(base_dir, relative_path)
        
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            continue
            
        print(f"Processing: {relative_path}")
        
        # Add filename as heading
        doc.add_heading(f'Файл: {relative_path}', level=1)
        
        # Add description
        doc.add_heading('Назначение:', level=2)
        p_desc = doc.add_paragraph(description)
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Add code
        doc.add_heading('Код:', level=2)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            # Add content
            paragraph = doc.add_paragraph(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Set font to monospaced
            for run in paragraph.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                
        except Exception as e:
            doc.add_paragraph(f"Ошибка при чтении файла: {str(e)}")

    # Ensure docs directory exists
    output_dir = os.path.join(base_dir, 'docs')
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_path = os.path.join(output_dir, 'Program_Listing.docx')
    doc.save(output_path)
    print(f"Listing saved to: {output_path}")

if __name__ == "__main__":
    create_program_listing()
